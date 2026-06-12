"""
模板脱敏工具 — 将模板中的具体信息替换为占位符。
用法: python template_anonymize.py <模板.docx> [--dry-run]
输出: 同目录下生成 <模板>.template.docx
"""
import sys
import os
import re
import shutil
from pathlib import Path

# 添加 docx skill 的 python-docx 路径
sys.path.insert(0, r"{{REASONIX_SKILLS}}docx\scripts")

from docx import Document


def replace_in_paragraph(para, replacements):
    """在段落的所有 run 中执行替换，跨 run 处理"""
    full_text = para.text
    new_text = full_text
    for old, new in replacements:
        new_text = new_text.replace(old, new)
    if new_text == full_text:
        return False

    # 清空所有 runs，写入新文本到第一个 run
    for i, run in enumerate(para.runs):
        if i == 0:
            run.text = new_text
        else:
            run.text = ""
    return True


def replace_in_table(table, replacements):
    """在表格所有单元格中执行替换"""
    changed = False
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if replace_in_paragraph(para, replacements):
                    changed = True
    return changed


def anonymize(input_path, replacements, output_path=None):
    """对 DOCX 执行脱敏替换"""
    if output_path is None:
        p = Path(input_path)
        output_path = str(p.parent / f"{p.stem}.template{p.suffix}")

    doc = Document(input_path)

    changed = 0
    for para in doc.paragraphs:
        if replace_in_paragraph(para, replacements):
            changed += 1
    for table in doc.tables:
        if replace_in_table(table, replacements):
            changed += 1

    doc.save(output_path)
    return changed


# ============================================================
# 各模板的替换映射
# ============================================================

# 诉状.docx 脱敏映射
诉状_替换 = [
    # --- 原告自然人信息 ---
    ("惠羿玮", "[姓名]"),
    ("，女，汉族", "，[性别]，[民族]"),
    ("1999年8月24日", "[出生年月日]"),
    ("411381199908240307", "[身份证号]"),
    ("住上海市普陀区千阳南路99弄19号", "住[户籍地]"),
    # --- 被告单位信息 ---
    ("漳州市龙文区安之妮汽车经营部", "[单位名称]"),
    ("统一社会信用代码：92350603MADUMPGH7P", "统一社会信用代码：[统一社会信用代码]"),
    ("经营者：郑阿妮", "经营者：[法定代表人]"),
    ("地址：福建省漳州市龙文区湘桥路2-18号", "地址：[住所地]"),
    ("联系电话：18205965540", "联系电话：[对方联系电话]"),
    # --- 合同/车辆 ---
    ("《二手（事故）车销售合同》", "[合同名称]"),
    ("车架号：LBV61AF09PS777048", "车架号：[车辆识别代号]"),
    # --- 金额 ---
    ("305,000元", "[购车款金额]元"),
    ("首付款61,000元", "首付款[首付款金额]元"),
    ("剩余款项244,000元", "剩余款项[融资金额]元"),
    ("赔偿金人民币915,000元", "赔偿金人民币[赔偿金金额]元"),
    ("合计人民币5,468元", "合计人民币[其他损失金额]元"),
    ("律师费人民币30000元", "律师费人民币[律师费金额]元"),
    # --- 法院 ---
    ("漳州市龙文区人民法院", "[法院名称]"),
    # --- 事实段落中的关键信息（尽可能保留叙事结构） ---
    ("原告因购车需求，于近期在被告经营场所看中一辆二手汽车", "原告因[购车目的]，于[看车时间]在被告经营场所看中一辆二手汽车"),
    ('被告故意隐瞒车辆真实车况，混淆\u201c事故\u201d程度', '被告故意隐瞒车辆真实车况，混淆\u201c事故\u201d程度'),
    ('原告自行委托第三方检测机构查验后发现，该车辆不仅存在严重调表行为，更系发生重大事故导致结构件多处切割、焊接的重大事故车', '[检测发现的事实描述]'),
    ('明知车辆存在重大事故及调表事实', '明知车辆存在[车辆瑕疵事实]'),
]

# 诉讼保全申请书.docx 脱敏映射
保全申请_替换 = [
    # --- 申请人 ---
    ("周红艳", "[姓名]"),
    ("，女，汉族", "，[性别]，[民族]"),
    ("1992年5月5日", "[出生年月日]"),
    ("住址安徽省南陵县籍山镇千峰村千山自然村30号", "住址[户籍地]"),
    # --- 被申请人 ---
    ("安徽狐狸小妖文化传播有限公司", "[单位名称]"),
    ("统一社会信用代码：91340111MA2TQ90P89", "统一社会信用代码：[统一社会信用代码]"),
    ("法定代表人：王佩佩", "法定代表人：[法定代表人]"),
    ("住所地：合肥市包河区马鞍山路与望江路交叉口绿地赢海大厦D座1307室", "住所地：[住所地]"),
    ("联系电话：18919684519（周波）、18756563307（都张超）", "联系电话：[对方联系电话]"),
    # --- 合同 ---
    ("《加盟合同2024版》", "[合同名称]"),
    ("于2024年4月30日签订", "于[合同签订日期]签订"),
    ("加盟费99,800元", "加盟费[加盟费金额]元"),
    ("软件系统及设备建设费60,000元", "软件系统及设备建设费[设备费金额]元"),
    ("货品保证金483,333.33元", "货品保证金[保证金金额]元"),
    # --- 金额 ---
    ("1,932,536.8元", "[诉讼标的额]元"),
    ("1,932,536.89元", "[诉讼标的额]元"),
    # --- 法院 ---
    ("合肥市包河区人民法院", "[法院名称]"),
    # --- 事实 ---
    ("申请人加盟被申请人经营的特许经营项目", "申请人加盟被申请人经营的[经营项目名称]"),
    ('被申请人授权使用的\u201c狐狸小妖护肤精选仓\u201d\u201c嘻选\u201d等商标均未依法注册', '被申请人授权使用的[授权商标清单]均未依法注册'),
    ("其在商务部备案的特许品牌在签约时也已过期", "其在商务部备案的特许品牌[备案情况]"),
    ("被申请人隐瞒上述事实，严重违反《商业特许经营管理条例》的如实告知义务", "被申请人隐瞒上述事实，严重违反[违反的法规名称]的如实告知义务"),
    ("被申请人经营状况不稳定", "被申请人[经营异常情况]"),
]

# ============================================================
# 主流程
# ============================================================

TEMPLATE_DIR = Path(r"{{REASONIX_SKILLS}}draft-legal-docs\templates")

FILES_MAP = {
    "诉状.docx": 诉状_替换,
    "诉讼保全申请书.docx": 保全申请_替换,
}


def main():
    dry_run = "--dry-run" in sys.argv

    for fname, replacements in FILES_MAP.items():
        input_path = TEMPLATE_DIR / fname
        if not input_path.exists():
            print(f"⚠ 跳过（不存在）: {fname}")
            continue

        output_path = TEMPLATE_DIR / f"{input_path.stem}.template{input_path.suffix}"

        if dry_run:
            print(f"\n{'='*60}")
            print(f"📄 [DRY-RUN] {fname} → {output_path.name}")
            print(f"   替换项: {len(replacements)} 条")
            for old, new in replacements:
                print(f"   {old!r} → {new!r}")
            continue

        print(f"\n{'='*60}")
        print(f"📄 处理: {fname}")

        changed = anonymize(str(input_path), replacements, str(output_path))
        print(f"   ✅ 修改 {changed} 个段落/单元格")
        print(f"   💾 输出: {output_path.name}")


if __name__ == "__main__":
    main()
