"""解析律所两个模板(案卷目录 + 办案小结)看结构"""
from docx import Document

def dump(path, label):
    print(f"\n{'='*60}")
    print(f"=== {label}: {path} ===")
    print('='*60)
    doc = Document(path)
    print(f"段落数: {len(doc.paragraphs)}")
    print(f"表格数: {len(doc.tables)}")
    print(f"节数: {len(doc.sections)}")
    # styles
    styles = set()
    for p in doc.paragraphs:
        if p.style and p.style.name:
            styles.add(p.style.name)
    print(f"段落样式: {sorted(styles)}")
    # sections
    for i, sec in enumerate(doc.sections):
        print(f"  Section {i}: page={sec.page_width}x{sec.page_height}, margins L/R/T/B={sec.left_margin}/{sec.right_margin}/{sec.top_margin}/{sec.bottom_margin}")

    # 所有段落(前 80 个)
    print(f"\n--- 前 80 段 ---")
    for i, p in enumerate(doc.paragraphs[:80]):
        text = p.text.strip()
        if not text:
            continue
        style = p.style.name if p.style else "?"
        print(f"  [{i:3d}] {style:20s} | {text[:80]}")

    # 所有表格
    print(f"\n--- 表格 ---")
    for ti, table in enumerate(doc.tables):
        print(f"  Table {ti}: {len(table.rows)} 行 x {len(table.columns)} 列")
        for ri, row in enumerate(table.rows[:25]):
            cells = " | ".join(c.text.strip()[:30] for c in row.cells)
            print(f"    R{ri}: {cells}")
        if len(table.rows) > 25:
            print(f"    ... 跳过 {len(table.rows) - 25} 行")

dump(r"{{HOME_WIN}}\.mavis\scratchpads\mvs_a2ae78a89f084128b874fe41fd47b238\template_catalog.docx", "律所案卷目录模板")
dump(r"{{HOME_WIN}}\.mavis\scratchpads\mvs_a2ae78a89f084128b874fe41fd47b238\template_summary.docx", "律所办案小结模板")
