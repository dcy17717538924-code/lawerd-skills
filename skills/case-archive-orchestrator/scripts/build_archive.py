"""build_archive.py - 试点:建 00-定卷/ + 复制改名 + 生成占位 docx"""
import os
import shutil
import subprocess
import sys
from pathlib import Path

# 路径
case_dir = Path(r"{{CASE_ROOT}}S施玲玲--------离婚")
archive_dir = case_dir / "00-定卷"
materials_dir = archive_dir / "入卷材料"

# 元数据(从 PDF 自动提取)
metadata = {
    "案件名": "施玲玲 vs 王一超 离婚纠纷",
    "案号": "(2025)沪0106民初31419号",
    "法院": "上海市静安区人民法院",
    "立案日期": "2025-07-08",
    "结案日期": "2025-07-22",
    "承办律师": "{{USER_FULL_NAME}}",
    "助理": "(空)",
    "收案日期": "2025-07-07",
    "年度": "2026",
    "归档日期": "2026-06-07",
    "是否有原件": "全复印件",
}

# 7 大类映射(2026-06-07 用户调整后版:9 个文件)
# 用户口述的 3 条规则(2026-06-07):
#   1. 1-3 一般在材料备份文件夹,以图片方式存在
#   2. 4-5 在案件目录或司法送达等文件夹
#   3. 6 是用户自己截图放在案件文件夹里,如果没有 skill 要提醒
# 用户决定:3 用 PDF 电子发票(不是图),Q2 暂时不放 2-诉讼准备/3-立案提交/ 的 docx
mappings = [
    # (大类序号, 新文件名, 原文件名, 大类名, 材料名, 注释)
    (1, "1委托合同.jpg",  "1委托合同.jpg",                            "审理结果", "委托合同",      "1-3 在 1-材料备份/, 图片"),
    (2, "2-1律师委托合同.jpg", "2-1律师委托合同.jpg",                    "委托手续", "律师委托合同",  "1-3 在 1-材料备份/, 图片"),
    (2, "2-2律师委托合同.jpg", "2-2律师委托合同.jpg",                    "委托手续", "律师委托合同",  "1-3 在 1-材料备份/, 图片"),
    (3, "3律师发函及律师费.pdf", "3律师发函及律师费.pdf",                "立案材料", "律师发函及律师费", "用户决定:用 PDF 电子发票,不是图"),
    (4, "4-1立案通知书_p001.pdf", "4-1立案通知书_p001.pdf",               "庭审材料", "立案通知书",   "4-5 在案件根,PDF"),
    (4, "4-2传票_p001.pdf",  "4-2传票_p001.pdf",                       "庭审材料", "传票",         "4-5 在案件根,PDF"),
    (4, "4-3举证通知书_p001.pdf", "4-3举证通知书_p001.pdf",               "庭审材料", "举证通知书",   "4-5 在案件根,PDF"),
    (5, "5民事调解书_p001.pdf", "5民事调解书_p001.pdf",                   "审理结果", "民事调解书",   "4-5 在案件根,PDF"),
    (6, "6散群截图.png",      "6散群截图.png",                            "签收归还", "散群截图",     "用户手动截图(规则3)"),
]

# 1. 建目录
archive_dir.mkdir(exist_ok=True)
materials_dir.mkdir(exist_ok=True)
print(f"已建: {archive_dir} + {materials_dir}")

# 2. 复制 + 改名(幂等:已存在则跳过,不动用户手动调好的)
print("\n=== 复制并改名(幂等模式) ===")
ok = []
skip = []
for cat, new_name, orig_name, cat_name, mat_name, *_ in mappings:
    src = case_dir / orig_name
    dst = materials_dir / new_name
    if not src.exists():
        print(f"  跳过(原文件不存在): {orig_name}")
        skip.append(new_name)
        continue
    if dst.exists():
        print(f"  跳过(已存在): {new_name}")
        ok.append(new_name)
        continue
    shutil.copy2(src, dst)
    print(f"  {cat}. {cat_name}: {new_name}  (从 {orig_name})")
    ok.append(new_name)

# 3. 装 python-docx
print("\n=== 装 python-docx ===")
r = subprocess.run(
    [sys.executable, "-m", "pip", "install", "--timeout", "30", "python-docx"],
    capture_output=True, text=True, timeout=60
)
print(r.stdout.split("\n")[-3] if r.stdout else "")
if r.returncode != 0:
    print("STDERR:", r.stderr[-300:])

# 4. 生成案卷目录.docx + 办案小结(草稿).docx
from docx import Document
from docx.shared import Pt

def create_docx(path, title, lines):
    """强制覆盖生成 docx(2026-06-07 试点最终版)"""
    if path.exists():
        path.unlink()  # 强制覆盖
        print(f"  覆盖: {path.name}")
    else:
        print(f"  生成: {path.name}")
    doc = Document()
    h = doc.add_heading(title, level=1)
    for line in lines:
        if line == "":
            doc.add_paragraph()
        else:
            p = doc.add_paragraph(line)
    doc.save(str(path))
    return True

print("\n=== 生成 docx ===")

# 元组格式(cat, new_name, orig_name, cat_name, mat_name, *extra) 末尾可能多 1 个 *extra 元素
def unpack_m(m):
    return m[0], m[1], m[2], m[3], m[4]

catalog_lines = [
    f"案号: {metadata['案号']}",
    f"案件名: {metadata['案件名']}",
    f"法院: {metadata['法院']}",
    f"承办律师: {metadata['承办律师']}",
    f"收案日期: {metadata['收案日期']}",
    f"立案日期: {metadata['立案日期']}",
    f"结案日期: {metadata['结案日期']}",
    f"归档日期: {metadata['归档日期']}",
    f"年度: {metadata['年度']}",
    f"是否有原件: {metadata['是否有原件']}",
    "",
    "=== 律所 7 大类 + 实际入卷材料(共 {} 件)===".format(len(ok)),
] + [
    f"  {unpack_m(m)[0]}. {unpack_m(m)[3]}: {unpack_m(m)[1]}" for m in mappings if unpack_m(m)[1] in ok
] + [
    "",
    "=== 律所目录 24 项占位(你后面填)===",
    "1. 审理结果: ____________________",
    "2. 委托手续: ____________________",
    "3. 立案材料: ____________________",
    "4. 庭审材料: ____________________",
    "5. 保全材料: (本案无)",
    "6. 签收及归还确认单: ____________________",
    "7. 其他: ____________________",
    "",
    "(后续 17 项细分, 按律所模板)",
]
create_docx(archive_dir / "案卷目录.docx", "案卷目录(占位)", catalog_lines)

summary_lines = [
    f"案件名: {metadata['案件名']}",
    f"案号: {metadata['案号']}",
    f"承办律师: {metadata['承办律师']}",
    f"结案日期: {metadata['结案日期']}",
    "",
    "一、案情摘要",
    "  原告施玲玲与被告王一超离婚纠纷,经上海市静安区人民法院主持调解,双方自愿达成协议。",
    "  原、被告于 2024 年 2 月 2 日登记结婚,婚后未生育子女。原告以夫妻感情破裂为由诉请离婚。",
    "",
    "二、争议焦点",
    "  1. 是否准予离婚",
    "  2. 夫妻共同财产分割(重点: 上海市静安区阳城路 283 弄 94 号 502 室房屋婚后共同还贷及增值部分)",
    "",
    "三、调解结果",
    "  1. 准予离婚",
    "  2. 被告 7 日内给付财产折价款 10,000 元",
    "  3. 被告 7 日内给付房屋婚后共同还贷及增值部分折价款 40,000 元",
    "  4. 案件受理费 100 元,原、被告各半负担(已履行)",
    "",
    "四、办案心得",
    "  (由 case-study-report skill 后续跑出,本文件为占位骨架)",
    "",
    "—— 本草稿由 case-archive-orchestrator 试点脚本自动生成, 待你/case-study-report 补全 ——",
]
create_docx(archive_dir / "办案小结(草稿).docx", "办案小结(草稿,占位)", summary_lines)

# 5. 报告
print("\n=== 试点完成 ===")
print(f"  00-定卷/ 路径: {archive_dir}")
print(f"  入卷材料/ 文件数: {len(ok)} (跳过 {len(skip)})")
print(f"  案卷目录.docx: {(archive_dir / '案卷目录.docx').exists()}")
print(f"  办案小结(草稿).docx: {(archive_dir / '办案小结(草稿).docx').exists()}")
