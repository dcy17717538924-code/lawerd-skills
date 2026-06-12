"""fill_templates.py - 复制律所模板 + 填字段(不重新建 docx)"""
import shutil
from pathlib import Path
from docx import Document

# 律所模板
TEMPLATE_DIR = Path(r"D:\wpsyunpan\229601413\WPS云盘\03 - 文书模板\申沪 - 入职律师培训文件(1)\入职律师培训文件\归档材料包")
TEMPLATE_CATALOG = TEMPLATE_DIR / "4 诉讼已结案案卷封面表-案卷目录.docx"
TEMPLATE_SUMMARY = TEMPLATE_DIR / "7{{USER_FIRM}}办案小结.docx"

# 案件归档
CASE_DIR = Path(r"{{CASE_ROOT}}S施玲玲--------离婚")
ARCHIVE_DIR = CASE_DIR / "00-定卷"
MATERIALS_DIR = ARCHIVE_DIR / "入卷材料"
CATALOG_OUT = ARCHIVE_DIR / "案卷目录.docx"
SUMMARY_OUT = ARCHIVE_DIR / "办案小结(草稿).docx"

# 元数据
metadata = {
    "案件名称": "施玲玲",
    "案件编号": "(2025)沪0106民初31419号",
    "承办律师": "{{USER_FULL_NAME}}",
    "业务类型": "民事",
    "收案日期": "2025-07-07",
    "结案日期": "2025-07-22",
    "归档日期": "2026-06-07",
    "是否有原件": "否",
    "年度": "2025",
    "审理结果": "调解",
    "卷内共计页数": "(需算 9 文件页数总和)",
    "审级": "一审",
    "保全情况": "无",
    "保全到期日": "无",
    "助理": "(空)",
}

# 7 大类中文字
CAT_NAMES = {1: "审理结果", 2: "委托手续", 3: "立案材料", 4: "庭审材料",
             5: "保全材料", 6: "签收及归还确认单", 7: "其他"}

# 我的 9 个材料(按入卷材料/ 顺序)
my_materials = [
    (1, "民事调解书",        "5民事调解书_p001.pdf",    "1-2", 1, "复印件"),
    (2, "律师委托合同",      "2-1律师委托合同.jpg",       "",    1, "图片"),
    (2, "律师委托合同",      "2-2律师委托合同.jpg",       "",    1, "图片"),
    (3, "律师发函及律师费",  "3律师发函及律师费.pdf",   "",    1, "电子发票"),
    (4, "立案通知书",        "4-1立案通知书_p001.pdf",    "",    1, ""),
    (4, "传票",              "4-2传票_p001.pdf",          "",    1, ""),
    (4, "举证通知书",        "4-3举证通知书_p001.pdf",    "",    1, ""),
    (4, "微信聊天记录",      "原 wechat_2025-07-22_182932_093.png", "", 1, "用户截图"),
    (6, "散群截图",          "6散群截图.png",            "",    1, "用户手动截图"),
]


def fill_catalog():
    """复制律所案卷目录模板 + 填字段"""
    if CATALOG_OUT.exists():
        CATALOG_OUT.unlink()
    shutil.copy2(TEMPLATE_CATALOG, CATALOG_OUT)
    doc = Document(str(CATALOG_OUT))
    table = doc.tables[0]
    # 律所表结构:R0-R6 是 7x4 顶部元数据(key-value pair 在列 0/1/2 与列 4/5/6)
    # 列布局: 0=label, 1=label(重复), 2=value, 3=label, 4=label(重复), 5=value, 6=label/value
    # 实际我数:每行 7 列,值在第 2 列(index 2)
    # R0: 案件名称 | 案件名称 | 薛社红 | 案件编号 | 无 | 无 | 无
    #   = (0=案件名称, 1=案件名称重复, 2=施玲玲, 3=案件编号, 4=无, 5=无, 6=无)
    # 我把第 2 列当 value,第 3+ 列当下一个 key-value
    # 但实际更稳:扫描每行,第一个匹配 key 的 cell 之后,下一 cell 是 value

    # 顶部 8 行(7 行元数据 + 1 行表头)
    # R0-R6 是元数据(key 在列 0,value 在列 2)
    # R7 是表头分隔("内容")
    # R8 是列头

    # 简化策略:R0-R6 每行 2 个 key-value 对
    # 左半:col 0=key, col 2=value
    # 右半:col 3=key, col 5=value(实际是 col 4,5,6 都填了)
    # 我用更稳的:每行只在 col 0 和 col 3 放 key,在 col 2 放 value
    # 然后再覆盖 col 5 等等

    # 元数据映射(只填主 key,key 在 col 0,value 在 col 2)
    top_keys = {
        "案件名称": metadata["案件名称"],
        "承办律师": metadata["承办律师"],
        "收案日期": metadata["收案日期"],
        "归档日期": metadata["归档日期"],
        "年度": metadata["年度"],
        "卷内共计页数": metadata["卷内共计页数"],
        "保全情况": metadata["保全情况"],
    }
    # col 3 的 key(我数了下:col 3 是 label)
    side_keys = {
        "案件编号": metadata["案件编号"],
        "业务类型": metadata["业务类型"],
        "结案日期": metadata["结案日期"],
        "是否有原件": metadata["是否有原件"],
        "审理结果": metadata["审理结果"],
        "审级": metadata["审级"],
        "保全到期日": metadata["保全到期日"],
    }

    # 顶部 7 行
    for ri in range(7):
        row = table.rows[ri]
        # 列 0 找 key → 列 2 填 value
        cell0 = row.cells[0].text.strip()
        if cell0 in top_keys:
            row.cells[2].text = top_keys[cell0]
        # 列 3 找 key → 列 5 填 value(律所 R0-R3 列 4/5/6 都填了"无",可能列 5 是主要 value)
        cell3 = row.cells[3].text.strip()
        if cell3 in side_keys:
            row.cells[5].text = side_keys[cell3]

    # R7 是表头分隔("内容"),不动
    # R8 是列头("序号 | | 材料名称 | 材料名称 | 页数 | 份数 | 备注"),不动

    # 材料清单(从 R9 开始,共 15 行,我的 9 个材料覆盖前 9 行,剩 6 行保留律所默认)
    # 列布局: 0=大类内序号, 1=大类, 2=材料名, 3=材料名(重复), 4=页数, 5=份数, 6=备注
    for ri, (cat, mat_name, fname, pages, copies, note) in enumerate(my_materials):
        if 9 + ri >= len(table.rows):
            break
        row = table.rows[9 + ri]
        # 大类内序号:本大类第几次出现(1, 1, 2, 1, 1, 1, 1, 1, 1)
        row.cells[0].text = str(_seq_in_cat(my_materials, ri, cat))
        row.cells[1].text = CAT_NAMES[cat]
        row.cells[2].text = mat_name
        row.cells[3].text = mat_name
        if pages:
            row.cells[4].text = pages
        if copies:
            row.cells[5].text = str(copies)
        if note:
            row.cells[6].text = note

    doc.save(str(CATALOG_OUT))


def _seq_in_cat(materials, idx, cat):
    """算第 idx 个材料在大类 cat 内是第几个"""
    cnt = 0
    for i in range(idx + 1):
        if materials[i][0] == cat:
            cnt += 1
    return cnt


def fill_summary():
    """复制律所办案小结模板 + 填字段"""
    if SUMMARY_OUT.exists():
        SUMMARY_OUT.unlink()
    shutil.copy2(TEMPLATE_SUMMARY, SUMMARY_OUT)
    doc = Document(str(SUMMARY_OUT))
    table = doc.tables[0]

    # 律所 R0-R3 4 行 4 列
    # R0: 案由 | 案由 | 案件名 | 案件名
    # R1: 委托人 | 委托人 | 受理机关 | 受理机关
    # R2: 简要案情 | 办案体会 | 简要案情 | 办案体会
    # R3: 处理结果 | 内容 | 内容 | 内容

    # 案由:离婚纠纷
    table.rows[0].cells[1].text = "离婚纠纷"
    table.rows[0].cells[2].text = "离婚纠纷"
    table.rows[0].cells[3].text = "离婚纠纷"

    # 委托人:施玲玲,受理机关:上海市静安区人民法院
    table.rows[1].cells[1].text = "施玲玲"
    table.rows[1].cells[2].text = "施玲玲"
    table.rows[1].cells[3].text = "上海市静安区人民法院"

    # 简要案情 + 办案体会
    brief = ("原告施玲玲与被告王一超离婚纠纷,经法院主持调解,双方自愿达成协议。"
             "原、被告 2024-02-02 登记结婚,婚后未生育子女,原告以夫妻感情破裂诉请离婚,被告同意。"
             "达成协议:1.准予离婚;2.被告 7 日内给付财产折价款 1 万元;"
             "3.被告 7 日内给付房屋婚后共同还贷及增值部分折价款 4 万元;"
             "4.案件受理费 100 元各半负担(已履行)。")
    experience = ("本案系协议离婚转诉讼离婚,争议焦点在房屋婚后共同还贷及增值部分的折价。"
                 "调解时强调「折价款」概念,以一次性 4 万元了结共有部分权益,避免评估周期拖延。"
                 "经验:离婚案件调解优先确认双方意愿一致性,争议在财产的案件可先固定「准予离婚」+"
                 "「一次性折价」思路,提升调解成功率。")
    table.rows[2].cells[1].text = brief
    table.rows[2].cells[2].text = brief
    table.rows[2].cells[3].text = experience

    # 处理结果
    table.rows[3].cells[1].text = "调解结案"
    table.rows[3].cells[2].text = "调解结案"
    table.rows[3].cells[3].text = "调解结案"

    # 段 4: 助理
    if len(doc.paragraphs) > 4:
        doc.paragraphs[4].text = f"承办律师助理:{metadata['助理']}"
    # 段 5: 日期
    if len(doc.paragraphs) > 5:
        doc.paragraphs[5].text = f"日期:{metadata['归档日期']}"

    doc.save(str(SUMMARY_OUT))


# === 主流程 ===

# 1. 备份律所模板到 00-定卷/_templates/(防律所路径变动)
ref_dir = ARCHIVE_DIR / "_templates"
ref_dir.mkdir(exist_ok=True)
for src, name in [(TEMPLATE_CATALOG, "template_catalog.docx"), (TEMPLATE_SUMMARY, "template_summary.docx")]:
    dst = ref_dir / name
    if not dst.exists():
        shutil.copy2(src, dst)
        print(f"备份模板: {dst}")

# 2. 填字段
print("\n=== 填案卷目录 ===")
fill_catalog()
print(f"  生成: {CATALOG_OUT}")

print("\n=== 填办案小结 ===")
fill_summary()
print(f"  生成: {SUMMARY_OUT}")

# 3. 报告
print("\n=== 完成 ===")
print(f"  案卷目录.docx:    {CATALOG_OUT}  ({CATALOG_OUT.stat().st_size} B)")
print(f"  办案小结(草稿).docx: {SUMMARY_OUT}  ({SUMMARY_OUT.stat().st_size} B)")
print(f"  律所模板备份:    {ref_dir}")
print(f"  入卷材料/(未动):  {MATERIALS_DIR}")
