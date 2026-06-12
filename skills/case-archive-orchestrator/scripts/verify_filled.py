"""验证 fill_templates 后的案卷目录 + 办案小结 跟律所同构"""
from docx import Document

print('=== 案卷目录.docx ===')
d = Document(r'{{CASE_ROOT}}S施玲玲--------离婚\00-定卷\案卷目录.docx')
print(f'段落: {len(d.paragraphs)}, 表格: {len(d.tables)}')
t = d.tables[0]
print(f'表格: {len(t.rows)} 行 x {len(t.columns)} 列')
print('顶部 7 行(元数据):')
for ri in range(7):
    row = t.rows[ri]
    cells = ' | '.join(c.text.strip()[:20] for c in row.cells)
    print(f'  R{ri}: {cells}')
print('R7-R8 表头:')
for ri in [7, 8]:
    row = t.rows[ri]
    cells = ' | '.join(c.text.strip()[:20] for c in row.cells)
    print(f'  R{ri}: {cells}')
print('材料清单(我填的 9 行 + 6 行律所默认):')
for ri in range(9, 24):
    row = t.rows[ri]
    cells = ' | '.join(c.text.strip()[:20] for c in row.cells)
    print(f'  R{ri}: {cells}')

print()
print('=== 办案小结(草稿).docx ===')
d = Document(r'{{CASE_ROOT}}S施玲玲--------离婚\00-定卷\办案小结(草稿).docx')
print(f'段落: {len(d.paragraphs)}, 表格: {len(d.tables)}')
t = d.tables[0]
print(f'表格: {len(t.rows)} 行 x {len(t.columns)} 列')
for ri, row in enumerate(t.rows):
    cells = ' | '.join(c.text.strip()[:35] for c in row.cells)
    print(f'  R{ri}: {cells}')
print('前 6 段:')
for pi, p in enumerate(d.paragraphs[:6]):
    if p.text.strip():
        print(f'  [{pi}] {p.text[:80]}')
