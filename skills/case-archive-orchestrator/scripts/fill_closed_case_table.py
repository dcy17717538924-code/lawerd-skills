"""Fill 诉讼案件已结案表.xlsx from archive metadata.

Default output:
  <CasePath>/00-定卷/诉讼案件已结案表.xlsx

The script first reads metadata from 案卷目录.docx when available, then applies
optional JSON/CLI overrides. Unknown fields are left blank or "无".
"""

from __future__ import annotations

import argparse
import json
import shutil
from pathlib import Path

from docx import Document
from openpyxl import load_workbook


SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_DIR = SCRIPT_DIR.parent
TEMPLATE = SKILL_DIR / "_references" / "templates" / "closed_case_table_template.xlsx"

COLUMNS = {
    "档案编号": 1,
    "委托日期": 2,
    "归档日期": 3,
    "承办律师": 4,
    "助理": 5,
    "案件类型": 6,
    "委托事项": 7,
    "案由": 8,
    "委托人": 9,
    "当事人": 10,
    "联系方式": 11,
    "诉讼地位": 12,
    "保全期限到期日": 13,
    "案件进度": 14,
    "待办事项": 15,
    "已付律师费": 16,
    "律师费阶段款": 17,
    "律师费风险金": 18,
    "对方": 19,
    "对方联系方式": 20,
    "受诉法院": 21,
    "法院案号": 22,
    "主办法官": 23,
    "法官电话": 24,
    "立案日期": 25,
    "开庭日期": 26,
    "备注1": 27,
    "结案日期": 28,
    "结案方式": 29,
    "裁决内容": 30,
    "阶段款是否支付": 31,
    "风险金是否支付": 32,
    "是否移交WordExcel文书电子档": 33,
    "是否移交纸质档扫描件": 34,
    "是否移交纸质档": 35,
    "代理人名字": 36,
    "散群时间": 37,
    "备注2": 38,
}


DEFAULTS = {
    "档案编号": "1",
    "承办律师": "{{USER_FULL_NAME}}",
    "助理": "",
    "案件类型": "民事",
    "委托事项": "一审",
    "联系方式": "无",
    "保全期限到期日": "无",
    "待办事项": "无",
    "律师费阶段款": "无",
    "律师费风险金": "无",
    "对方联系方式": "无",
    "主办法官": "无",
    "法官电话": "无",
    "备注1": "无",
    "阶段款是否支付": "/",
    "风险金是否支付": "/",
    "是否移交WordExcel文书电子档": "是",
    "是否移交纸质档扫描件": "是",
    "是否移交纸质档": "是",
    "散群时间": "已散群",
    "备注2": "无",
}


def read_catalog_metadata(case_path: Path) -> dict[str, str]:
    catalog = case_path / "00-定卷" / "案卷目录.docx"
    if not catalog.exists():
        return {}

    doc = Document(str(catalog))
    if not doc.tables:
        return {}
    table = doc.tables[0]
    data: dict[str, str] = {}
    mapping = {
        "案件名称": "案件名称",
        "承办律师": "承办律师",
        "收案日期": "委托日期",
        "归档日期": "归档日期",
        "案件编号": "法院案号",
        "业务类型": "案件类型",
        "结案日期": "结案日期",
        "审理结果": "结案方式",
        "保全到期日": "保全期限到期日",
    }
    for row in table.rows[:7]:
        cells = row.cells
        pairs = []
        if len(cells) > 2:
            pairs.append((cells[0].text.strip(), cells[2].text.strip()))
        if len(cells) > 5:
            pairs.append((cells[3].text.strip(), cells[5].text.strip()))
        for key, value in pairs:
            target = mapping.get(key)
            if target and value:
                data[target] = value
    if "案件名称" in data:
        data.setdefault("案由", "")
        data.setdefault("当事人", data["案件名称"])
        data.setdefault("委托人", data["案件名称"].split("与")[0].split("诉")[0])
    data.setdefault("代理人名字", data.get("承办律师", "{{USER_FULL_NAME}}"))
    data.setdefault("案件进度", data.get("结案方式", "已结案"))
    return data


def load_overrides(path: Path | None) -> dict[str, str]:
    if not path:
        return {}
    with path.open("r", encoding="utf-8") as f:
        obj = json.load(f)
    return {str(k): "" if v is None else str(v) for k, v in obj.items()}


def fill_table(case_path: Path, overrides: dict[str, str]) -> Path:
    archive_dir = case_path / "00-定卷"
    archive_dir.mkdir(exist_ok=True)
    output = archive_dir / "诉讼案件已结案表.xlsx"

    shutil.copy2(TEMPLATE, output)
    wb = load_workbook(output)
    ws = wb.active

    data = dict(DEFAULTS)
    data.update(read_catalog_metadata(case_path))
    data.update(overrides)

    for key, col in COLUMNS.items():
        ws.cell(2, col).value = data.get(key, "")

    wb.save(output)
    print(f"生成: {output}")
    for key in ["委托日期", "归档日期", "承办律师", "案件类型", "案由", "委托人", "当事人", "受诉法院", "法院案号", "结案日期", "结案方式", "散群时间"]:
        print(f"  {key}: {data.get(key, '')}")
    return output


def main() -> None:
    parser = argparse.ArgumentParser(description="填写诉讼案件已结案表")
    parser.add_argument("-CasePath", required=True, help="案件根目录")
    parser.add_argument("--metadata-json", help="UTF-8 JSON，字段名使用结案表中文列名")
    parser.add_argument("--set", action="append", default=[], help="覆盖字段，格式：字段=值，可重复")
    args = parser.parse_args()

    overrides = load_overrides(Path(args.metadata_json)) if args.metadata_json else {}
    for item in args.set:
        if "=" not in item:
            raise ValueError(f"--set 必须是 字段=值 格式: {item}")
        key, value = item.split("=", 1)
        overrides[key] = value

    fill_table(Path(args.CasePath), overrides)


if __name__ == "__main__":
    main()
